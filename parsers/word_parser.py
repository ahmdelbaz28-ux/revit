"""
word_parser.py — FireAI Word Specification Parser
Parses project specifications from Word documents.

Extracts:
    - Project title/identifier
    - Floor information
    - Ceiling specifications
    - Special requirements/notes
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, field

logger = logging.getLogger("fireai.word_parser")


# ═══════════════════════════════════════════════════════
# DATA CLASS
# ═══════════════════════════════════════════════════════

@dataclass
class WordParseResult:
    """Result of parsing Word document."""
    source_file: str
    success: bool
    title: str = ""
    project_name: str = ""
    floor: str = ""
    building: str = ""
    ceiling_specs: List[Dict] = field(default_factory=list)
    requirements: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


# ═══════════════════════════════════════════════════════
# WORD PARSER
# ═══════════════════════════════════════════════════════

class WordParser:
    """
    Parses Word documents for project specifications.
    
    USAGE:
        parser = WordParser()
        result = parser.parse("project_specs.docx")
        
        if result.success:
            print(f"Project: {result.project_name}")
            print(f"Floor: {result.floor}")
    """

    # Patterns for extraction
    FLOOR_PATTERNS = [
        r'floor\s*(\d+)',
        r'level\s*(\d+)',
        r'Floor\s*(\d+)',
        r'Level\s*(\d+)',
        r'(\d+)st\s*floor',
        r'(\d+)nd\s*floor',
        r'(\d+)rd\s*floor',
        r'(\d+)th\s*floor',
    ]
    
    BUILDING_PATTERNS = [
        r'building\s*([A-Z])',
        r'tower\s*([A-Z])',
        r'block\s*([A-Z])',
        r'([A-Z])\s*building',
    ]
    
    CEILING_PATTERNS = [
        r'ceiling\s*height[:\s]*(\d+\.?\d*)\s*m',
        r'height[:\s]*(\d+\.?\d*)\s*m',
        r'flat\s*ceiling[:\s]*(\d+\.?\d*)',
        r'suspended\s*ceiling[:\s]*(\d+\.?\d*)',
    ]

    def __init__(self):
        pass

    def parse(self, file_path: str) -> WordParseResult:
        """
        Parse Word document.

        Args:
            file_path: Path to .docx file. MUST be under
                FIREAI_ALLOWED_UPLOAD_DIRS (V124 security hardening).

        Returns:
            WordParseResult with extracted info
        """
        # V126: Path security + file-size cap
        from parsers._path_security import (
            UnsafePathError,
            validate_input_path,
            validate_file_size,
        )
        _ALLOWED_EXTENSIONS = frozenset({".docx", ".doc"})
        _MAX_FILE_SIZE_BYTES = int(os.getenv("FIREAI_WORD_MAX_FILE_SIZE_BYTES", 25 * 1024 * 1024))  # 25 MB default
        try:
            safe_path = validate_input_path(
                file_path,
                allowed_extensions=_ALLOWED_EXTENSIONS,
                parser_name="WordParser",
            )
            validate_file_size(
                safe_path,
                max_size_bytes=_MAX_FILE_SIZE_BYTES,
                parser_name="WordParser",
            )
        except FileNotFoundError as e:
            return WordParseResult(source_file=file_path, success=False, errors=[str(e)])
        except UnsafePathError as e:
            return WordParseResult(source_file=file_path, success=False, errors=[f"SECURITY: {e}"])

        file_path = str(safe_path)
        result = WordParseResult(source_file=file_path, success=False)

        try:
            from docx import Document
            
            doc = Document(str(safe_path))
            
            # Extract from all paragraphs
            all_text = '\n'.join(p.text for p in doc.paragraphs)
            
            # Extract title (first heading)
            result.title = self._extract_title(doc.paragraphs)
            
            # Extract project info
            result.project_name = self._extract_project_name(all_text)
            result.floor = self._extract_floor(all_text)
            result.building = self._extract_building(all_text)
            
            # Extract ceiling specs
            result.ceiling_specs = self._extract_ceiling_specs(all_text)
            
            # Extract requirements
            result.requirements = self._extract_requirements(doc.paragraphs)
            
            # Extract notes
            result.notes = self._extract_notes(doc.paragraphs)
            
            result.success = bool(result.title or result.project_name or result.floor)
            
        except ImportError as e:
            result.errors.append(f"Missing dependency: python-docx not installed")
        except Exception as e:
            result.errors.append(f"Parse error: {type(e).__name__}: {e}")
            
        return result

    def _extract_title(self, paragraphs) -> str:
        """Extract title from first heading."""
        for para in paragraphs:
            if para.style.name.startswith('Heading'):
                return para.text.strip()
        return ""

    def _extract_project_name(self, text: str) -> str:
        """Extract project/building name."""
        # Look for project patterns
        patterns = [
            r'Project[:\s]*([^\n]+)',
            r'Building[:\s]*([^\n]+)',
            r'Tower\s*([A-Z])',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
                
        return ""

    def _extract_floor(self, text: str) -> str:
        """Extract floor number."""
        for pattern in self.FLOOR_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return f"Floor {match.group(1)}"
        return ""

    def _extract_building(self, text: str) -> str:
        """Extract building/block identifier."""
        for pattern in self.BUILDING_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return f"Building {match.group(1)}"
        return ""

    def _extract_ceiling_specs(self, text: str) -> List[Dict]:
        """Extract ceiling specifications."""
        specs = []
        
        for pattern in self.CEILING_PATTERNS:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    height = float(match.group(1))
                    specs.append({
                        'type': 'flat',
                        'height_m': height,
                    })
                except (ValueError, IndexError):
                    continue
                    
        return specs

    def _extract_requirements(self, paragraphs) -> List[str]:
        """Extract requirement bullets."""
        requirements = []
        
        for para in paragraphs:
            text = para.text.strip()
            
            # Check for bullet points
            if text.startswith('•') or text.startswith('- ') or text.startswith('* '):
                clean_text = text.lstrip('•-* ').strip()
                
                # Filter relevant requirements
                if any(kw in clean_text.lower() for kw in [
                    'detector', 'alarm', 'fire', 'sprinkler', 
                    'system', 'zone', 'coverage', 'code'
                ]):
                    requirements.append(clean_text)
                    
        return requirements

    def _extract_notes(self, paragraphs) -> List[str]:
        """Extract notes section."""
        notes = []
        in_notes = False
        
        for para in paragraphs:
            text = para.text.strip()
            
            # Check for notes section
            if 'note' in text.lower() and len(text) < 20:
                in_notes = True
                continue
                
            if in_notes and text:
                notes.append(text)
                
        return notes[:10]  # Limit to 10 notes


# ═══════════════════════════════════════════════════════
# CONVENIENCE FUNCTION
# ═══════════════════════════════════════════════════════

def parse_word(file_path: str) -> WordParseResult:
    """Quick parse Word file."""
    parser = WordParser()
    return parser.parse(file_path)